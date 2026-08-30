"""Builders for synthetic test documents.

Everything the test suite and the evaluation cases read is generated here, in
every format the system accepts. The data is invented: the names, numbers and
addresses belong to nobody, and the Aadhaar-style numbers are constructed to
carry a valid checksum only so the validation path can be exercised.

Generating rather than committing binaries keeps the fixtures inspectable and
lets an evaluation case vary one field at a time, which is exactly what a
discrepancy test needs.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document as DocxDocument
from openpyxl import Workbook
from PIL import Image, ImageDraw, ImageFont
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas


def make_pdf(path: str | Path, title: str, fields: dict[str, str], *, pages: int = 1) -> Path:
    """A text-layer PDF laid out as label/value lines."""
    path = Path(path)
    path.parent.mkdir(parents=True, exist_ok=True)
    pdf = canvas.Canvas(str(path), pagesize=A4)
    width, height = A4

    items = list(fields.items())
    per_page = max(1, -(-len(items) // pages))
    for page_index in range(pages):
        pdf.setFont("Helvetica-Bold", 16)
        pdf.drawString(70, height - 80, title)
        pdf.setFont("Helvetica", 11)
        y = height - 120
        for label, value in items[page_index * per_page : (page_index + 1) * per_page]:
            pdf.drawString(70, y, f"{label}: {value}")
            y -= 22
        pdf.setFont("Helvetica-Oblique", 8)
        pdf.drawString(70, 50, f"Page {page_index + 1} of {pages}")
        pdf.showPage()
    pdf.save()
    return path


def make_scanned_pdf(path: str | Path, title: str, fields: dict[str, str]) -> Path:
    """A PDF with no text layer: an image of the text, as a scanner produces."""
    import fitz

    image_path = Path(path).with_suffix(".tmp.png")
    make_image(image_path, title, fields)

    path = Path(path)
    document = fitz.open()
    pixmap = fitz.Pixmap(str(image_path))
    page = document.new_page(width=pixmap.width, height=pixmap.height)
    page.insert_image(fitz.Rect(0, 0, pixmap.width, pixmap.height), filename=str(image_path))
    document.save(str(path))
    document.close()
    image_path.unlink(missing_ok=True)
    return path


def make_image(
    path: str | Path,
    title: str,
    fields: dict[str, str],
    *,
    size: tuple[int, int] = (1200, 900),
    blur: bool = False,
) -> Path:
    path = Path(path)
    path.parent.mkdir(parents=True, exist_ok=True)
    image = Image.new("RGB", size, "white")
    draw = ImageDraw.Draw(image)
    scale = max(1.0, size[0] / 1200)
    heading = _font(int(34 * scale))
    body = _font(int(26 * scale))
    draw.text((60, 50), title, fill="black", font=heading)
    y = 120
    for label, value in fields.items():
        draw.text((60, y), f"{label}: {value}", fill="black", font=body)
        y += int(46 * scale)
    if blur:
        from PIL import ImageFilter

        image = image.filter(ImageFilter.GaussianBlur(radius=6))
    image.save(path)
    return path


def make_docx(path: str | Path, title: str, fields: dict[str, str]) -> Path:
    path = Path(path)
    path.parent.mkdir(parents=True, exist_ok=True)
    document = DocxDocument()
    document.add_heading(title, level=1)
    document.add_paragraph("Submitted for credit assessment.")

    table = document.add_table(rows=0, cols=2)
    for label, value in fields.items():
        row = table.add_row().cells
        row[0].text = label
        row[1].text = str(value)
    document.save(str(path))
    return path


def make_xlsx(path: str | Path, sheets: dict[str, list[list[str]]]) -> Path:
    path = Path(path)
    path.parent.mkdir(parents=True, exist_ok=True)
    workbook = Workbook()
    workbook.remove(workbook.active)
    for name, rows in sheets.items():
        sheet = workbook.create_sheet(title=name[:31])
        for row in rows:
            sheet.append(row)
    workbook.save(str(path))
    return path


def _font(size: int):
    """A scalable face, falling back to the bitmap default on older Pillow."""
    try:
        return ImageFont.load_default(size=size)
    except TypeError:
        return ImageFont.load_default()


# --------------------------------------------------------------------------
# Land title documents
# --------------------------------------------------------------------------
# A chain of title is a sequence, not a single document, so these are built as
# a set: each deed hands the property to the buyer named in the next one. The
# helper exists so a test can break exactly one link and change nothing else,
# which is what makes a chain test mean something.
def make_sale_deed(
    path: str | Path,
    *,
    seller: str,
    buyer: str,
    registered_on: str,
    survey_number: str = "42/1B",
    consideration: str = "Rs. 42,00,000",
    registration_number: str = "BLR-4-2024-1187",
) -> Path:
    return make_pdf(
        path,
        "DEED OF ABSOLUTE SALE",
        {
            "Seller": seller,
            "Buyer": buyer,
            "Survey Number": survey_number,
            "Property Address": "Site 14, Ward 8, Kengeri Hobli, Bengaluru 560060",
            "Sale Consideration": consideration,
            "Registration Number": registration_number,
            "Registration Date": registered_on,
        },
    )


def make_encumbrance_certificate(
    path: str | Path,
    *,
    owner: str,
    period_from: str = "01/01/2013",
    period_to: str = "01/07/2026",
    encumbrance: str = "NIL",
    survey_number: str = "42/1B",
) -> Path:
    return make_pdf(
        path,
        "ENCUMBRANCE CERTIFICATE - OFFICE OF THE SUB-REGISTRAR",
        {
            "Owner Name": owner,
            "Survey Number": survey_number,
            "Property Address": "Site 14, Ward 8, Kengeri Hobli, Bengaluru 560060",
            "Period From": period_from,
            "Period To": period_to,
            "Encumbrance": encumbrance,
            "Registration Number": "EC-BLR-2026-9921",
        },
    )


def make_khata_certificate(
    path: str | Path, *, owner: str, survey_number: str = "42/1B"
) -> Path:
    return make_pdf(
        path,
        "KHATA CERTIFICATE - BRUHAT BENGALURU MAHANAGARA PALIKE",
        {
            "Owner Name": owner,
            "Khata Number": "184/14/8/42-1B",
            "Survey Number": survey_number,
            "Property Address": "Site 14, Ward 8, Kengeri Hobli, Bengaluru 560060",
            "Assessment Year": "2026-27",
        },
    )


def make_property_tax_receipt(
    path: str | Path, *, owner: str, receipt_date: str = "10/05/2026"
) -> Path:
    return make_pdf(
        path,
        "PROPERTY TAX PAID RECEIPT",
        {
            "Owner Name": owner,
            "Khata Number": "184/14/8/42-1B",
            "Property Address": "Site 14, Ward 8, Kengeri Hobli, Bengaluru 560060",
            "Assessment Year": "2026-27",
            "Amount Paid": "Rs. 8,420",
            "Receipt Date": receipt_date,
        },
    )


def make_land_pack(
    directory: str | Path,
    *,
    applicant: str = "Ravi Kumar",
    chain: tuple[tuple[str, str, str], ...] = (
        ("Anil Sharma", "Meera Reddy", "14/03/2015"),
        ("Meera Reddy", "Suresh Kumar", "02/09/2019"),
        ("Suresh Kumar", "Ravi Kumar", "21/06/2024"),
    ),
    encumbrance: str = "NIL",
    khata_owner: str | None = None,
) -> dict[str, Path]:
    """A complete land pack: the deeds, the EC, the khata and the tax receipt.

    ``chain`` is the sequence of transfers. Passing a tuple whose seller does
    not match the previous buyer is how a test produces a broken title without
    disturbing anything else.
    """
    directory = Path(directory)
    documents: dict[str, Path] = {}
    for index, (seller, buyer, registered) in enumerate(chain, start=1):
        documents[f"deed_{index}"] = make_sale_deed(
            directory / f"SaleDeed{registered[-4:]}_{index}.pdf",
            seller=seller,
            buyer=buyer,
            registered_on=registered,
            registration_number=f"BLR-4-{registered[-4:]}-{1100 + index}",
        )
    owner = chain[-1][1] if chain else applicant
    documents["ec"] = make_encumbrance_certificate(
        directory / "EncumbranceCertificate.pdf", owner=owner, encumbrance=encumbrance
    )
    documents["khata"] = make_khata_certificate(
        directory / "KhataCertificate.pdf", owner=khata_owner or owner
    )
    documents["tax"] = make_property_tax_receipt(
        directory / "PropertyTaxReceipt.pdf", owner=khata_owner or owner
    )
    return documents


def make_encumbrance_certificate_with_table(
    path: str | Path,
    *,
    transactions: list[tuple[str, str, str, str]],
    period_from: str = "01/01/2013",
    period_to: str = "01/07/2026",
    survey_number: str = "42/1B",
) -> Path:
    """An EC as one actually looks: a header block over a table of entries.

    ``transactions`` is a list of (date, nature, executant, claimant). The
    label/value builder cannot express this, and a certificate's whole content
    is the table — reading it as a summary was the limitation this fixture
    exists to close.
    """
    path = Path(path)
    path.parent.mkdir(parents=True, exist_ok=True)
    pdf = canvas.Canvas(str(path), pagesize=A4)
    width, height = A4

    pdf.setFont("Helvetica-Bold", 14)
    pdf.drawString(50, height - 60, "ENCUMBRANCE CERTIFICATE")
    pdf.setFont("Helvetica", 10)
    pdf.drawString(50, height - 78, "OFFICE OF THE SUB-REGISTRAR, KENGERI, BENGALURU")

    y = height - 110
    for label, value in (
        ("Survey Number", survey_number),
        ("Property Address", "Site 14, Ward 8, Kengeri Hobli, Bengaluru 560060"),
        ("Period From", period_from),
        ("Period To", period_to),
    ):
        pdf.setFont("Helvetica", 10)
        pdf.drawString(50, y, f"{label}: {value}")
        y -= 16

    y -= 14
    pdf.setFont("Helvetica-Bold", 9)
    columns = (50, 105, 205, 330, 440)
    for text, x in zip(
        ("Sl.", "Date", "Nature", "Executant", "Claimant"), columns, strict=True
    ):
        pdf.drawString(x, y, text)
    pdf.line(50, y - 4, 545, y - 4)
    y -= 20

    pdf.setFont("Helvetica", 9)
    for index, (date, nature, executant, claimant) in enumerate(transactions, start=1):
        for text, x in zip(
            (str(index), date, nature, executant, claimant), columns, strict=True
        ):
            pdf.drawString(x, y, text)
        y -= 16

    y -= 18
    pdf.setFont("Helvetica-Oblique", 9)
    pdf.drawString(
        50, y, f"{len(transactions)} entries found for the period stated above."
    )
    pdf.save()
    return path
