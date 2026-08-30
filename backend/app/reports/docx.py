"""DOCX rendering.

The renderer walks the template's section list and prints whatever the
generator produced for each one. It contains no knowledge of banking and makes
no decisions: if a value is not in the report JSON, it does not appear on the
page.
"""

from __future__ import annotations

import io
from typing import Any

from docx import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

SEVERITY_COLOURS = {
    "HIGH": RGBColor(0xB3, 0x26, 0x1E),
    "MEDIUM": RGBColor(0xB5, 0x6A, 0x00),
    "LOW": RGBColor(0x4A, 0x4A, 0x4A),
}


def render_docx(report: dict[str, Any], template: dict[str, Any]) -> bytes:
    document = DocxDocument()

    document.add_heading(report.get("title", "Document Verification Report"), level=0)
    if report.get("subtitle"):
        subtitle = document.add_paragraph(report["subtitle"])
        subtitle.alignment = WD_ALIGN_PARAGRAPH.LEFT
        subtitle.runs[0].italic = True

    for section in template.get("sections", []):
        key = section.get("key")
        content = report.get(key)
        if content in (None, [], {}):
            # An empty section is still worth naming, so the reader can tell
            # "nothing found" apart from "not checked".
            document.add_heading(section.get("title", key), level=1)
            document.add_paragraph("Nothing to report in this section.")
            continue

        document.add_heading(section.get("title", key), level=1)
        section_type = section.get("type", "object")

        if section_type == "findings":
            _render_findings(document, content)
        elif section_type == "profile":
            _render_profile(document, content)
        elif section_type == "table":
            _render_table(document, section, content)
        else:
            _render_object(document, section, content)

    if report.get("disclaimer"):
        document.add_page_break()
        document.add_heading("Important", level=2)
        note = document.add_paragraph(report["disclaimer"])
        note.runs[0].font.size = Pt(9)

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _render_object(document, section: dict[str, Any], content: dict[str, Any]) -> None:  # noqa: ANN001
    labels = section.get("fields") or {}
    for key, value in content.items():
        if key == "counts":
            continue
        label = labels.get(key, key.replace("_", " ").title())
        if isinstance(value, list):
            document.add_paragraph(f"{label}:")
            for item in value:
                document.add_paragraph(str(item), style="List Bullet")
        else:
            paragraph = document.add_paragraph()
            paragraph.add_run(f"{label}: ").bold = True
            paragraph.add_run(str(value))


def _render_table(document, section: dict[str, Any], rows: list[dict[str, Any]]) -> None:  # noqa: ANN001
    columns = section.get("columns") or []
    if not columns or not rows:
        return

    table = document.add_table(rows=1, cols=len(columns))
    table.style = "Light Grid Accent 1"
    header = table.rows[0].cells
    for index, column in enumerate(columns):
        header[index].text = column.get("title", column["key"])

    for row in rows:
        cells = table.add_row().cells
        for index, column in enumerate(columns):
            cells[index].text = str(row.get(column["key"], ""))


def _render_profile(document, content: dict[str, Any]) -> None:  # noqa: ANN001
    table = document.add_table(rows=1, cols=4)
    table.style = "Light Grid Accent 1"
    for index, title in enumerate(["Field", "Value", "Status", "Sources"]):
        table.rows[0].cells[index].text = title

    for name, field in content.items():
        cells = table.add_row().cells
        cells[0].text = name.replace("_", " ").title()
        value = field.get("value", "")
        candidates = field.get("candidates") or []
        # A conflict is printed as a conflict, not as a single chosen value.
        cells[1].text = " / ".join(candidates) if candidates else str(value)
        cells[2].text = str(field.get("status", ""))
        cells[3].text = "; ".join(
            f"{source['document']} p{source['page']}" for source in field.get("sources", [])
        )


def _render_findings(document, rows: list[dict[str, Any]]) -> None:  # noqa: ANN001
    """Findings get prose blocks, not table cells.

    A reviewer has to be able to read the two conflicting values, where each
    came from, and what to do about it, without decoding a wide table.
    """
    for finding in rows:
        heading = document.add_paragraph()
        run = heading.add_run(
            f"{finding['id']} — {finding['severity']} — "
            f"{finding['type'].replace('_', ' ').title()}"
        )
        run.bold = True
        colour = SEVERITY_COLOURS.get(str(finding.get("severity")))
        if colour is not None:
            run.font.color.rgb = colour

        if finding.get("field") and finding["field"] != "NOT_AVAILABLE":
            document.add_paragraph(f"Field: {finding['field']}")

        for side in ("1", "2"):
            source = finding.get(f"document_{side}")
            if source and source != "NOT_AVAILABLE":
                page = finding.get(f"page_{side}") or 0
                page_text = f", page {page}" if page else ""
                document.add_paragraph(
                    f"{source}{page_text}: {finding.get(f'value_{side}', '')}",
                    style="List Bullet",
                )

        if finding.get("explanation"):
            document.add_paragraph(finding["explanation"])
        document.add_paragraph(
            f"Status: {finding.get('status')} | Confidence: {finding.get('confidence')} | "
            f"Evidence verified: {'yes' if finding.get('verified') else 'no'} | "
            f"Rule: {finding.get('rule_id')}"
        )
        action = document.add_paragraph()
        action.add_run("Recommended action: ").bold = True
        action.add_run(str(finding.get("recommended_action", "")))
        document.add_paragraph()
